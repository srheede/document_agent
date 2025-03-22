import os
import sys
import json
import asyncio
import logging
import re
import zipfile
import io
import traceback
from pathlib import Path
from datetime import datetime, timezone

import requests
from dotenv import load_dotenv
from dataclasses import dataclass
from typing import List, Dict, Any

from PIL import Image
import pytesseract
import PyPDF2
import fitz  # PyMuPDF
from bs4 import BeautifulSoup
from docx import Document

# Load environment variables
load_dotenv()

# Initialize clients
from openai import AsyncOpenAI
from supabase import create_client, Client

openai_client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))
supabase: Client = create_client(os.getenv("SUPABASE_URL"), os.getenv("SUPABASE_SERVICE_KEY"))

logging.basicConfig(level=logging.INFO)

@dataclass
class ProcessedChunk:
    file_path: str
    chunk_number: int
    title: str
    summary: str
    content: str
    metadata: Dict[str, Any]
    embedding: List[float]

def chunk_text(text: str, chunk_size: int = 5000) -> List[str]:
    chunks = []
    start = 0
    text_length = len(text)
    while start < text_length:
        end = start + chunk_size
        if end >= text_length:
            chunks.append(text[start:].strip())
            break
        chunk = text[start:end]
        code_block = chunk.rfind('```')
        if code_block != -1 and code_block > chunk_size * 0.3:
            end = start + code_block
        elif "\n\n" in chunk:
            last_break = chunk.rfind("\n\n")
            if last_break > chunk_size * 0.3:
                end = start + last_break
        elif ". " in chunk:
            last_period = chunk.rfind(". ")
            if last_period > chunk_size * 0.3:
                end = start + last_period + 1
        current_chunk = text[start:end].strip()
        if current_chunk:
            chunks.append(current_chunk)
        start = max(start + 1, end)
    return chunks

async def get_title_and_summary(chunk: str, file_path: str) -> Dict[str, str]:
    system_prompt = (
        "You are an AI that extracts titles and summaries from documentation chunks.\n"
        "Return a JSON object with 'title' and 'summary' keys.\n"
        "For the title: If this seems like the start of a document, extract its title. "
        "If it's a middle chunk, derive a descriptive title.\n"
        "For the summary: Create a concise summary of the main points in this chunk."
    )
    try:
        response = await openai_client.chat.completions.create(
            model=os.getenv("LLM_MODEL", "gpt-4o-mini"),
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"File: {file_path}\n\nContent:\n{chunk[:1000]}..."}
            ],
            response_format={"type": "json_object"}
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        logging.error(f"Error getting title and summary: {e}")
        return {"title": "Error processing title", "summary": "Error processing summary"}

async def get_embedding(text: str) -> List[float]:
    try:
        response = await openai_client.embeddings.create(
            model="text-embedding-3-small",
            input=text
        )
        return response.data[0].embedding
    except Exception as e:
        logging.error(f"Error getting embedding: {e}")
        return [0] * 1536

async def process_chunk(chunk: str, chunk_number: int, file_path: str) -> ProcessedChunk:
    extracted = await get_title_and_summary(chunk, file_path)
    embedding = await get_embedding(chunk)
    metadata = {
        "source": "local_documents",
        "chunk_size": len(chunk),
        "processed_at": datetime.now(timezone.utc).isoformat(),
        "file_path": file_path
    }
    return ProcessedChunk(
        file_path=file_path,
        chunk_number=chunk_number,
        title=extracted.get("title", ""),
        summary=extracted.get("summary", ""),
        content=chunk,
        metadata=metadata,
        embedding=embedding
    )

async def insert_chunk(chunk: ProcessedChunk):
    data = {
        "url": chunk.file_path,
        "chunk_number": chunk.chunk_number,
        "title": chunk.title,
        "summary": chunk.summary,
        "content": chunk.content,
        "metadata": chunk.metadata,
        "embedding": chunk.embedding
    }
    try:
        result = supabase.table("site_pages").insert(data).execute()
        logging.info(f"Inserted chunk {chunk.chunk_number} for {chunk.file_path}")
        return result
    except Exception as e:
        logging.error(f"Error inserting chunk: {e}")
        return None

async def process_and_store_document(file_path: str, text: str):
    chunks = chunk_text(text)
    tasks = [process_chunk(chunk, i, file_path) for i, chunk in enumerate(chunks)]
    processed_chunks = await asyncio.gather(*tasks)
    insert_tasks = [insert_chunk(chunk) for chunk in processed_chunks]
    await asyncio.gather(*insert_tasks)

# --- Text extraction functions ---

def extract_text_from_image(image: Image.Image) -> str:
    return pytesseract.image_to_string(image)

def extract_text_from_docx_images(docx_path: str) -> str:
    extracted_text = ""
    try:
        with zipfile.ZipFile(docx_path) as z:
            for file_name in z.namelist():
                if file_name.startswith("word/media/") and file_name.lower().endswith(('.png', '.jpg', '.jpeg')):
                    with z.open(file_name) as img_file:
                        image = Image.open(io.BytesIO(img_file.read()))
                        extracted_text += extract_text_from_image(image) + "\n"
    except Exception as e:
        logging.error(f"Error extracting images from DOCX {docx_path}: {e}")
    return extracted_text

def extract_text_from_shapes(doc):
    texts = []
    for textbox in doc.element.xpath('//w:txbxContent'):
        for p in textbox.xpath('.//w:p'):
            texts.append("".join(t.text for t in p.xpath('.//w:t') if t.text))
    return "\n".join(texts)

def doc_to_text(doc_path: str, temp_folder: str = "/tmp/temp") -> str:
    try:
        Path(temp_folder).mkdir(parents=True, exist_ok=True)
        doc = Document(doc_path)
        texts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(cell.text)
        for section in doc.sections:
            for p in section.header.paragraphs:
                texts.append(p.text)
            for p in section.footer.paragraphs:
                texts.append(p.text)
        shape_text = extract_text_from_shapes(doc)
        if shape_text:
            texts.append(shape_text)
        full_text = "\n".join(t.strip() for t in texts if t.strip())
        if not full_text.strip():
            logging.info("No text found in DOCX; falling back to OCR on images.")
            full_text = extract_text_from_docx_images(doc_path)
        return full_text
    except Exception as e:
        logging.error(f"Error in doc_to_text for {doc_path}: {e}", exc_info=True)
        raise

def extract_text_from_pdf_images(pdf_path: str) -> str:
    extracted_text = ""
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            pix = page.get_pixmap()
            image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            extracted_text += extract_text_from_image(image) + "\n"
    except Exception as e:
        logging.error(f"Error extracting images from PDF {pdf_path}: {e}")
    return extracted_text

def pdf_to_text(pdf_path: str) -> str:
    try:
        with open(pdf_path, "rb") as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file)
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages)
        if not text.strip():
            logging.info("No text extracted via PyPDF2; falling back to OCR using PyMuPDF.")
            text = extract_text_from_pdf_images(pdf_path)
        return text
    except Exception as e:
        logging.error(f"Error in pdf_to_text for {pdf_path}: {e}", exc_info=True)
        raise Exception(f"Failed to extract text from PDF {pdf_path}: {e}") from e

def convert_file_to_text(file_path: str) -> str:
    try:
        ext = Path(file_path).suffix.lower()
        if ext == ".docx":
            return doc_to_text(file_path)
        elif ext == ".pdf":
            return pdf_to_text(file_path)
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        elif ext == ".rtf":
            from striprtf.striprtf import rtf_to_text
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                raw_text = f.read()
            return rtf_to_text(raw_text).replace("|", "")
        elif ext in [".html", ".htm"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                html = f.read()
            soup = BeautifulSoup(html, "html.parser")
            return soup.get_text()
        elif ext in [".png", ".jpg", ".jpeg", ".bmp"]:
            image = Image.open(file_path)
            return extract_text_from_image(image)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    except Exception as e:
        logging.error(f"Error converting {file_path} to text: {e}", exc_info=True)
        raise

def clean_text(text: str) -> str:
    try:
        return re.sub(r"[^\x09\x0A\x0D\x20-\x7E]+", " ", text).strip()
    except Exception as e:
        logging.error(f"Error cleaning text: {e}", exc_info=True)
        raise

async def process_document(file_path: Path):
    try:
        logging.info(f"Processing {file_path}")
        text = convert_file_to_text(str(file_path))
        text = clean_text(text)
        if text:
            full_file_path = str(file_path.resolve())
            await process_and_store_document(full_file_path, text)
        else:
            logging.info(f"No text extracted from {file_path}")
    except Exception as e:
        logging.error(f"Error processing document {file_path}: {e}", exc_info=True)

async def main():
    if len(sys.argv) < 2:
        print("Usage: python store_documents.py <folder_path>")
        return
    folder = Path(sys.argv[1])
    if not folder.is_dir():
        print("Provided path is not a directory.")
        return
    files = list(folder.rglob("*.*"))
    supported_ext = {".docx", ".pdf", ".txt", ".rtf", ".html", ".htm", ".png", ".jpg", ".jpeg", ".bmp"}
    files_to_process = [f for f in files if f.suffix.lower() in supported_ext]
    if not files_to_process:
        logging.info("No supported files found in the folder.")
        return
    await asyncio.gather(*(process_document(f) for f in files_to_process))

if __name__ == "__main__":
    asyncio.run(main())
